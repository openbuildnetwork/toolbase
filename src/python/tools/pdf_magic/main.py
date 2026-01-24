import io
import sys
from pypdf import PdfReader, PdfWriter
from PIL import Image
import fitz
from docx import Document
from docx.shared import Pt


# Helper function to render page to image
def render_page_to_image(pdf_bytes, page_num, dpi=150):
    """
    Render a PDF page to a PIL Image using PyMuPDF.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num]

    # Render page to pixmap at specified DPI
    mat = fitz.Matrix(dpi / 72, dpi / 72)  # 72 is default DPI
    pix = page.get_pixmap(matrix=mat)

    # Convert pixmap to PIL Image
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    doc.close()

    return img


def compress_pdf_to_images(file_bytes, level="recommended"):
    """
    Compress PDF by rendering pages to images and rebuilding.
    This achieves the same result as the TypeScript implementation.
    """
    try:
        # Configuration based on levels
        if level == "extreme":
            quality = 30
            dpi = 100  # Lower DPI for smaller file
        elif level == "recommended":
            quality = 60
            dpi = 150  # Medium DPI
        else:  # less
            quality = 85
            dpi = 200  # Higher DPI for better quality

        # Load PDF
        input_stream = io.BytesIO(bytes(file_bytes))
        reader = PdfReader(input_stream)
        num_pages = len(reader.pages)

        # Create new PDF writer
        writer = PdfWriter()

        # Process each page
        for page_num in range(num_pages):
            # Render page to image
            img = render_page_to_image(file_bytes, page_num, dpi)

            # Compress image
            img_bytes = io.BytesIO()
            if img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            img.save(img_bytes, format="JPEG", quality=quality, optimize=True)
            img_bytes.seek(0)

            # We need to create a page with the image
            # PyPDF doesn't have direct image-to-page, so we'll use a different approach

            # Create a simple PDF page with the image using reportlab-like approach
            # Actually, let's use a simpler method: create a minimal PDF structure

            # For now, let's use the img2pdf approach if available in Pyodide
            # If not, we'll need to construct PDF manually

            # Since img2pdf might not be available, let's use PyMuPDF to create pages
            pass

        # Alternative: Use PyMuPDF for the entire process
        return compress_with_pymupdf(file_bytes, quality, dpi)

    except Exception as e:
        return {"error": str(e)}


def compress_with_pymupdf(file_bytes, quality=60, dpi=150):
    """
    Complete compression using PyMuPDF (fitz).
    This is the most efficient approach in Python.
    """
    try:
        # Open source PDF
        src_doc = fitz.open(stream=file_bytes, filetype="pdf")

        # Create new PDF
        new_doc = fitz.open()

        for page_num in range(len(src_doc)):
            # Render page to pixmap
            page = src_doc[page_num]
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat)

            # Convert to PIL Image for JPEG compression
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Compress as JPEG
            img_bytes = io.BytesIO()
            if img.mode in ("RGBA", "P", "LA"):
                img = img.convert("RGB")
            img.save(img_bytes, format="JPEG", quality=quality, optimize=True)
            img_bytes.seek(0)

            # Create new page with compressed image
            img_doc = fitz.open(stream=img_bytes.getvalue(), filetype="jpeg")
            pdf_bytes = img_doc.convert_to_pdf()
            img_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
            new_doc.insert_pdf(img_pdf)

            img_doc.close()
            img_pdf.close()

        # Save result
        output = io.BytesIO()
        new_doc.save(output)
        result = output.getvalue()

        src_doc.close()
        new_doc.close()

        if len(result) == 0:
            raise Exception("Compression resulted in empty file")

        return list(result)

    except Exception as e:
        return {"error": str(e)}


def compress_pdf_content(file_bytes, level="recommended"):
    """
    Main compression function.
    Uses PyMuPDF for image-based compression.
    """
    try:
        if level == "extreme":
            return compress_with_pymupdf(file_bytes, quality=30, dpi=100)
        elif level == "recommended":
            return compress_with_pymupdf(file_bytes, quality=60, dpi=150)
        else:  # less
            return compress_with_pymupdf(file_bytes, quality=85, dpi=200)
    except Exception as e:
        return {"error": str(e)}


def protect_pdf(file_bytes, password, owner_password=None, permissions=None):
    """
    Encrypt PDF with a password using PyMuPDF.
    """
    try:
        # Open source PDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        # Default encryption options
        encryption_method = fitz.PDF_ENCRYPT_AES_256

        # Prepare permissions bitmask
        # By default, we might want to restrict significantly if permissions are provided
        # PDF Permissions Bitmask:
        # bit 3: print
        # bit 4: modify
        # bit 5: copy
        # bit 6: annotate
        # bit 9: fill forms
        # bit 10: accessibility
        # bit 11: document assembly

        perm_mask = 0
        if permissions:
            if permissions.get("printing"):
                perm_mask |= fitz.PDF_PERM_PRINT
            if permissions.get("modifying"):
                perm_mask |= fitz.PDF_PERM_MODIFY
            if permissions.get("copying"):
                perm_mask |= fitz.PDF_PERM_COPY
            if permissions.get("annotating"):
                perm_mask |= fitz.PDF_PERM_ANNOTATE
            if permissions.get("fillingForms"):
                perm_mask |= fitz.PDF_PERM_FORM
            # Always allow accessibility for compliance if not specified
            if permissions.get("accessibility", True):
                perm_mask |= 1 << 9
        else:
            # If no permissions object, allow everything
            perm_mask = -1

        # Save results to memory
        output = io.BytesIO()

        doc.save(
            output,
            encryption=encryption_method,
            user_pw=password,
            owner_pw=owner_password or password,
            permissions=perm_mask,
        )

        result = output.getvalue()
        doc.close()

        return list(result)
    except Exception as e:
        return {"error": str(e)}


def detect_pdf_elements(file_bytes, page_index):
    """
    Detect text blocks and images on a specific page.
    Returns a list of elements with their coordinates and page dimensions.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page = doc[page_index]

        elements = []
        pW, pH = page.rect.width, page.rect.height

        # Get text blocks
        blocks = page.get_text("blocks")
        for b in blocks:
            # b = (x0, y0, x1, y1, text, block_no, block_type)
            # block_type 0 is text, 1 is image
            if b[6] == 0:
                elements.append(
                    {
                        "id": f"orig-text-{b[5]}",
                        "type": "text",
                        "content": b[4],
                        "rect": [b[0], b[1], b[2], b[3]],
                        "existing": True,
                    }
                )

        # Get image locations
        image_list = page.get_image_info(xrefs=True)
        for i, img in enumerate(image_list):
            # img has bbox (x0, y0, x1, y1)
            elements.append(
                {
                    "id": f"orig-img-{img['number']}",
                    "type": "image",
                    "rect": img["bbox"],
                    "existing": True,
                }
            )

        doc.close()
        return {"elements": elements, "width": pW, "height": pH}
    except Exception as e:
        return {"error": str(e)}


def apply_pdf_edits(file_bytes, edits):
    """
    Apply a list of edits (additions and redactions) to a PDF.
    Expects percentages for x, y, width, height (0-100).
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        # Group edits by page
        page_edits = {}
        for edit in edits:
            p_idx = edit.get("pageIndex", 0)
            if p_idx not in page_edits:
                page_edits[p_idx] = []
            page_edits[p_idx].append(edit)

        for p_idx, p_list in page_edits.items():
            if p_idx >= len(doc):
                continue
            page = doc[p_idx]
            pW, pH = page.rect.width, page.rect.height

            # Helper to convert percentage to points
            def to_pdf_rect(el):
                w = (el["width"] / 100) * pW
                h = (el["height"] / 100) * pH
                x = (el["x"] / 100) * pW - w / 2
                y = (el["y"] / 100) * pH - h / 2
                return fitz.Rect(x, y, x + w, y + h)

            def hex_to_rgb(hex_str):
                if not hex_str:
                    return (0, 0, 0)
                hex_str = hex_str.lstrip("#")
                return tuple(int(hex_str[i : i + 2], 16) / 255.0 for i in (0, 2, 4))

            # 1. Handle Existing Content Removal (Redactions)
            # If an existing element was moved or deleted, it should have redact=True or old coordinates
            for edit in p_list:
                if edit.get("existing") and (edit.get("redact") or edit.get("moved")):
                    # Use originalRect if available
                    orig = edit.get("originalRect")
                    if orig:
                        page.add_redact_annot(orig, fill=(1, 1, 1))
                elif edit["type"] == "whiteout":
                    page.add_redact_annot(to_pdf_rect(edit), fill=(1, 1, 1))

            page.apply_redactions()

            # 2. Handle New Additions or Repositioned Content
            for edit in p_list:
                if edit.get("redact"):
                    continue  # Skip if only for redaction

                # If existing but not moved, skip re-adding
                if edit.get("existing") and not edit.get("moved"):
                    continue

                rect = to_pdf_rect(edit)
                color = hex_to_rgb(edit.get("color", "#000000"))

                if edit["type"] == "text":
                    page.insert_textbox(
                        rect,
                        edit["content"],
                        fontsize=edit.get("fontSize", 12),
                        color=color,
                    )
                elif edit["type"] in ["image", "drawing"]:
                    import base64

                    try:
                        img_data = edit["content"].split(",")[1]
                        img_bytes = base64.b64decode(img_data)
                        page.insert_image(rect, stream=img_bytes)
                    except:
                        pass
                elif edit["type"] == "shape":
                    if edit.get("shapeType") == "rectangle":
                        page.draw_rect(
                            rect, color=color, width=edit.get("strokeWidth", 1)
                        )
                    elif edit.get("shapeType") == "circle":
                        page.draw_oval(
                            rect, color=color, width=edit.get("strokeWidth", 1)
                        )

        output = io.BytesIO()
        doc.save(output)
        result = output.getvalue()
        doc.close()
        return list(result)
    except Exception as e:
        return {"error": str(e)}


def pdf_to_word(file_bytes):
    """
    Convert PDF to Word document (.docx) using PyMuPDF and python-docx.

    """
    try:
        # Open source PDF
        pdf_data = bytes(file_bytes)
        doc = fitz.open(stream=pdf_data, filetype="pdf")

        # Create new Word document
        word_doc = Document()

        for page in doc:
            elements = []

            # 1. Extract Text Blocks
            # Using get_text("blocks") is generally more robust for text extraction than get_text("dict")
            # because it groups text into logical blocks, often corresponding to paragraphs,
            # which is better for reconstructing document flow in a Word document.
            # get_text("dict") provides more granular span-level details but can be harder
            # to reassemble into coherent paragraphs without complex layout analysis.
            text_blocks = page.get_text("blocks")
            for b in text_blocks:
                # b: (x0, y0, x1, y1, text, block_no, block_type)
                if b[6] == 0:  # Text
                    elements.append(
                        {
                            "type": "text",
                            "content": b[4].strip(),
                            "bbox": (b[0], b[1], b[2], b[3]),
                            "y": b[1],  # Sort by top edge
                            "x": b[0],
                        }
                    )

            # 2. Extract Images
            # get_image_info(xrefs=True) is more robust for images than get_text("dict")
            # because it directly queries the PDF's image objects (XObjects) and inline images,
            # providing reliable bounding boxes and xrefs to extract the raw image data.
            # get_text("dict") might miss some images or provide less direct access to their data.
            image_info = page.get_image_info(xrefs=True)
            for img in image_info:
                bbox = img["bbox"]
                # Skip images with invalid bounding boxes or zero dimensions
                if not bbox or bbox[2] - bbox[0] <= 0 or bbox[3] - bbox[1] <= 0:
                    continue
                elements.append(
                    {
                        "type": "image",
                        "xref": img["xref"],
                        "bbox": bbox,
                        "y": bbox[1],
                        "x": bbox[0],
                        "width": bbox[2] - bbox[0],
                        "height": bbox[3] - bbox[1],
                    }
                )

            # 3. Sort Layout
            # Primary: Vertical (y), Secondary: Horizontal (x)
            elements.sort(key=lambda e: (e["y"], e["x"]))

            # 4. Construct Document
            for el in elements:
                if el["type"] == "text":
                    text = el["content"]
                    if text:
                        # Clean up text (handling hyphens etc could happen here)
                        word_doc.add_paragraph(text)

                elif el["type"] == "image":
                    try:
                        # Recover image binary
                        xref = el["xref"]
                        if xref == 0:
                            # Inline image, get_image_info should provide the image data directly
                            # if it's an inline image. If xref is 0, it means it's an inline image
                            # and its data should be in img['image'] if available.
                            # For simplicity, we'll skip if xref is 0 and no direct image data is easily accessible
                            # via xref. PyMuPDF's get_image_info usually provides xrefs for all extractable images.
                            continue

                        pix = fitz.Pixmap(doc, xref)

                        # Fix colorspace if CMYK or other non-RGB
                        if pix.n - pix.alpha > 3:  # Check if it's not RGB or RGBA
                            pix = fitz.Pixmap(fitz.csRGB, pix)  # Convert to RGB

                        # Save to PNG format via buffer
                        img_bytes = pix.tobytes("png")
                        img_stream = io.BytesIO(img_bytes)

                        # Scale to fit page (approximate 1pt = 1/72 inch)
                        # An 80% scale usually looks better within margins
                        w_inch = Inches(el["width"] / 72 * 0.9)

                        word_doc.add_picture(img_stream, width=w_inch)
                        pix = None  # Release pixmap resources

                    except Exception as img_err:
                        # Use a placeholder or skip if image fails
                        # print(f"Image error: {img_err}")
                        pass

            # Add page break
            if page.number < len(doc) - 1:
                word_doc.add_page_break()

        # Finalize
        output = io.BytesIO()
        word_doc.save(output)
        result = output.getvalue()

        doc.close()
        return list(result)

    except Exception as e:
        return {"error": str(e)}


def handle_request(action, data):
    if action == "compress":
        file_bytes = data.get("file_bytes")
        level = data.get("level", "recommended")
        return compress_pdf_content(file_bytes, level)
    elif action == "protect":
        file_bytes = data.get("file_bytes")
        password = data.get("password")
        owner_password = data.get("owner_password")
        permissions = data.get("permissions")
        return protect_pdf(file_bytes, password, owner_password, permissions)
    elif action == "detect":
        file_bytes = data.get("file_bytes")
        page_index = data.get("page_index", 0)
        return detect_pdf_elements(file_bytes, page_index)
    elif action == "apply_edits":
        file_bytes = data.get("file_bytes")
        edits = data.get("edits", [])
        return apply_pdf_edits(file_bytes, edits)
    elif action == "pdf_to_word":
        file_bytes = data.get("file_bytes")
        return pdf_to_word(file_bytes)

    return {"error": f"Unknown action: {action}"}

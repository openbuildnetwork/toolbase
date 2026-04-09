import io
import fitz
from PIL import Image
# Lazy imports for optional features

def pdf_to_word(file_bytes):
    try:
        from docx import Document
        from docx.shared import Inches
        doc = fitz.open(stream=bytes(file_bytes), filetype="pdf")
        word_doc = Document()

        for page in doc:
            # 1. Extract Text
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0])) # Sort by vertical position (y), then x
            
            for b in blocks:
                if b[6] == 0:  # Text block
                    text = b[4].strip()
                    if text:
                        word_doc.add_paragraph(text)

            # 2. Extract Images
            image_list = page.get_image_info(xrefs=True)
            for img in image_list:
                xref = img["xref"]
                if xref == 0: continue # Skip inline images for now
                
                try:
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n - pix.alpha > 3:
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    
                    img_data = pix.tobytes("png")
                    
                    # Add to Word (scaled down slightly to fit margins)
                    # Calculating width in inches based on PDF point width
                    img_width_pts = img["bbox"][2] - img["bbox"][0]
                    word_doc.add_picture(io.BytesIO(img_data), width=Inches(img_width_pts / 72 * 0.9))
                except:
                    pass

            word_doc.add_page_break()

        output = io.BytesIO()
        word_doc.save(output)
        return list(output.getvalue())

    except Exception as e:
        return {"error": str(e)}

def pdf_to_images(file_bytes, dpi=150, img_format="JPEG"):
    try:
        file_bytes = bytes(file_bytes)
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        images = []
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(dpi/72, dpi/72))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            out = io.BytesIO()
            img.save(out, format=img_format, quality=85)
            images.append(out.getvalue())
        return images
    except Exception as e:
        return {"error": str(e)}

def images_to_pdf(files_bytes_list):
    try:
        doc = fitz.open()
        for img_bytes in files_bytes_list:
            img_doc = fitz.open(stream=bytes(img_bytes), filetype="img")
            pdf_bytes = img_doc.convert_to_pdf()
            img_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
            doc.insert_pdf(img_pdf)
            img_doc.close()
            img_pdf.close()
            
        output = io.BytesIO()
        doc.save(output)
        return list(output.getvalue())
    except Exception as e:
        return {"error": str(e)}

def html_to_pdf(html_bytes):
    try:
        from xhtml2pdf import pisa
        html_content = html_bytes.decode('utf-8')
        output = io.BytesIO()
        pisa_status = pisa.CreatePDF(io.StringIO(html_content), dest=output)
        if pisa_status.err:
            return {"error": "PDF generation failed"}
        return list(output.getvalue())
    except Exception as e:
        return {"error": str(e)}